"""File parsing service: extract raw text from various file formats."""

import logging
import re
import unicodedata
from pathlib import Path

logger = logging.getLogger(__name__)


class FileParsingError(Exception):
    """Raised when a file cannot be parsed."""
    pass


def extract_text(file_path: Path, file_type: str) -> str:
    """Extract raw text from a file.

    Args:
        file_path: Path to the file.
        file_type: One of 'txt', 'md', 'docx', 'pdf'.

    Returns:
        Extracted text content.

    Raises:
        FileParsingError: If the file cannot be parsed.
    """
    if not file_path.exists():
        raise FileParsingError(f"File not found: {file_path}")

    extractors = {
        "txt": _extract_txt,
        "md": _extract_md,
        "markdown": _extract_md,
        "docx": _extract_docx,
        "pdf": _extract_pdf,
    }

    extractor = extractors.get(file_type.lower())
    if not extractor:
        raise FileParsingError(f"Unsupported file type: {file_type}")

    try:
        text = extractor(file_path)
    except FileParsingError:
        raise
    except Exception as e:
        raise FileParsingError(f"Failed to parse {file_path}: {e}") from e

    text = _postprocess(text)

    if not text.strip():
        raise FileParsingError(f"No text content found in {file_path}")

    return text


def _extract_txt(path: Path) -> str:
    """Extract text from a plain text file."""
    encodings = ["utf-8-sig", "utf-8", "gbk", "gb2312", "latin-1"]
    for enc in encodings:
        try:
            return path.read_text(encoding=enc)
        except (UnicodeDecodeError, UnicodeError):
            continue
    raise FileParsingError(f"Cannot decode file with supported encodings: {path}")


def _extract_md(path: Path) -> str:
    """Extract text from a Markdown file, stripping formatting syntax."""
    raw = _extract_txt(path)

    # Remove HTML tags
    text = re.sub(r"<[^>]+>", "", raw)
    # Remove images
    text = re.sub(r"!\[([^\]]*)\]\([^)]+\)", r"\1", text)
    # Convert links to just their text
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    # Remove bold/italic markers (keep the text)
    text = re.sub(r"\*{1,3}([^*]+)\*{1,3}", r"\1", text)
    text = re.sub(r"_{1,3}([^_]+)_{1,3}", r"\1", text)
    # Remove strikethrough
    text = re.sub(r"~~([^~]+)~~", r"\1", text)
    # Remove inline code backticks
    text = re.sub(r"`([^`]+)`", r"\1", text)
    # Convert headings to plain text (remove # markers)
    text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)
    # Remove horizontal rules
    text = re.sub(r"^[-*_]{3,}\s*$", "", text, flags=re.MULTILINE)
    # Remove blockquote markers
    text = re.sub(r"^>\s?", "", text, flags=re.MULTILINE)

    return text


def _extract_docx(path: Path) -> str:
    """Extract text from a DOCX file."""
    try:
        from docx import Document
    except ImportError:
        raise FileParsingError("python-docx is required for DOCX parsing. Install with: pip install python-docx")

    doc = Document(str(path))
    paragraphs = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    return "\n\n".join(paragraphs)


def _extract_pdf(path: Path) -> str:
    """Extract text from a PDF file."""
    try:
        import pdfplumber
    except ImportError:
        raise FileParsingError("pdfplumber is required for PDF parsing. Install with: pip install pdfplumber")

    pages_text = []

    with pdfplumber.open(str(path)) as pdf:
        if len(pdf.pages) == 0:
            raise FileParsingError(f"PDF has no pages: {path}")

        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages_text.append(text)

    if not pages_text:
        raise FileParsingError(f"Could not extract text from PDF (may be scanned/image-based): {path}")

    return "\n\n".join(pages_text)


def _postprocess(text: str) -> str:
    """Normalize whitespace and unicode characters."""
    # Normalize unicode: smart quotes to straight quotes, preserve em-dashes
    text = unicodedata.normalize("NFC", text)
    text = text.replace("\u2018", "'").replace("\u2019", "'")  # smart single quotes
    text = text.replace("\u201c", '"').replace("\u201d", '"')  # smart double quotes
    # Preserve em-dash (\u2014) and en-dash (\u2013)

    # Collapse multiple blank lines (max 2 consecutive)
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Strip trailing whitespace per line
    text = re.sub(r"[ \t]+$", "", text, flags=re.MULTILINE)
    # Strip leading/trailing whitespace from entire text
    text = text.strip()

    return text


def detect_file_type(filename: str) -> str:
    """Detect file type from extension."""
    ext = Path(filename).suffix.lower()
    type_map = {
        ".txt": "txt",
        ".md": "md",
        ".markdown": "md",
        ".docx": "docx",
        ".pdf": "pdf",
    }
    file_type = type_map.get(ext)
    if not file_type:
        raise FileParsingError(f"Unsupported file extension: {ext}")
    return file_type


def count_words(text: str) -> int:
    """Estimate word count, handling both CJK and Latin text."""
    # Count CJK characters individually
    cjk_count = len(re.findall(r"[\u4e00-\u9fff\u3400-\u4dbf\uf900-\ufaff]", text))
    # Count Latin words
    latin_words = re.findall(r"[a-zA-Z0-9]+(?:'[a-zA-Z]+)?", text)
    return cjk_count + len(latin_words)
